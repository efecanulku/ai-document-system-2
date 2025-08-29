import os
import json
import io
from typing import Optional
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document as DocxDocument
import openpyxl
from sqlalchemy.orm import Session

from app.models.document import Document, DocumentChunk
from app.services.gemini_service import GeminiService

class DocumentProcessor:
    def __init__(self):
        self.gemini_service = GeminiService()
        # Tesseract path'ini ayarla (Windows için)
        tesseract_cmd = os.getenv("TESSERACT_CMD", "tesseract")
        if os.name == 'nt' and tesseract_cmd == "tesseract":
            # Windows için varsayılan path
            possible_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
            ]
            for path in possible_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break
    
    async def process_document(self, document: Document, db: Session) -> bool:
        """Dökümanı işle ve AI analizi yap"""
        try:
            print(f"🚀 Döküman işleme başlıyor: {document.filename}")
            print(f"📁 Dosya yolu: {document.file_path}")
            print(f"🔧 Dosya tipi: {document.file_type}")
            
            # İçeriği çıkar
            print(f"📖 Metin çıkarma başlıyor...")
            content_text = await self.extract_text_content(document.file_path, document.file_type)
            
            if not content_text:
                print(f"❌ No content extracted from {document.filename}")
                return False
            
            print(f"✅ Metin çıkarma tamamlandı: {len(content_text)} karakter")
            
            # AI analizi yap
            print(f"🤖 AI analizi başlıyor...")
            try:
                print(f"📝 Özet oluşturuluyor...")
                summary = await self.gemini_service.generate_summary(content_text)
                print(f"✅ Özet tamamlandı: {len(summary)} karakter")
                
                print(f"🔑 Anahtar kelimeler çıkarılıyor...")
                keywords = await self.gemini_service.extract_keywords(content_text)
                print(f"✅ Anahtar kelimeler tamamlandı: {len(keywords)} kelime")
                
                print(f"🧠 Embedding oluşturuluyor...")
                embeddings = await self.gemini_service.generate_embeddings(content_text)
                print(f"✅ Embedding tamamlandı: {len(embeddings)} vektör")
                
            except Exception as ai_error:
                print(f"❌ AI analizi hatası: {ai_error}")
                print(f"❌ AI hata türü: {type(ai_error)}")
                # AI hatası durumunda varsayılan değerler kullan
                summary = "AI analizi başarısız"
                keywords = ["hata"]
                embeddings = []
            
            # Veritabanını güncelle
            print(f"💾 Veritabanı güncelleniyor...")
            document.content_text = content_text
            document.summary = summary
            document.keywords = json.dumps(keywords, ensure_ascii=False)
            document.embeddings = json.dumps(embeddings)
            document.processed = True
            
            print(f"✅ Veritabanı güncellendi")
            
            # Dökümanı parçalara böl ve kaydet
            print(f"✂️ Döküman chunk'lara bölünüyor...")
            await self.create_document_chunks(document, content_text, db)
            
            print(f"💾 Veritabanı commit ediliyor...")
            db.commit()
            print(f"✅ Döküman işleme tamamlandı: {document.filename}")
            return True
            
        except Exception as e:
            print(f"❌ Document processing error for {document.filename}: {e}")
            print(f"❌ Hata türü: {type(e)}")
            print(f"❌ Hata detayı: {str(e)}")
            import traceback
            print(f"❌ Traceback: {traceback.format_exc()}")
            db.rollback()
            return False
    
    async def extract_text_content(self, file_path: str, file_type: str) -> Optional[str]:
        """Dosyadan metin içeriği çıkar - Akıllı format tespiti ile"""
        try:
            # Dosya header'ını kontrol et ve gerçek formatı tespit et
            real_file_type = await self._detect_real_file_type(file_path, file_type)
            print(f"🔍 Dosya uzantısı: {file_type}, Gerçek format: {real_file_type}")
            
            if real_file_type == "pdf":
                return await self._extract_pdf_text(file_path)
            elif real_file_type == "docx":
                return await self._extract_docx_text(file_path)
            elif real_file_type == "doc":
                return await self._extract_doc_text(file_path)
            elif real_file_type in ["xlsx", "xls"]:
                return await self._extract_excel_text(file_path)
            elif real_file_type in ["jpg", "jpeg", "png", "gif", "bmp", "tiff"]:
                return await self._extract_image_text(file_path)
            elif real_file_type in ["html", "htm"]:
                return await self._extract_html_text(file_path)
            elif real_file_type == "txt":
                return await self._extract_txt_text(file_path)
            else:
                print(f"⚠️ Desteklenmeyen dosya formatı: {real_file_type}")
                return None
        except Exception as e:
            print(f"Text extraction error: {e}")
            return None
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """PDF'den metin çıkar - PyMuPDF ile optimize edilmiş"""
        text = ""
        try:
            pdf_document = fitz.open(file_path)
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
                else:
                    # Eğer metin yoksa OCR dene
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    try:
                        ocr_text = pytesseract.image_to_string(Image.open(io.BytesIO(img_data)), lang='tur+eng')
                        text += ocr_text + "\n"
                    except:
                        pass
            pdf_document.close()
        except Exception as e:
            print(f"PDF extraction error: {e}")
        
        # Metni temizle ve normalize et
        cleaned_text = self._clean_and_normalize_text(text)
        return cleaned_text
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """DOCX dosyasından metin çıkar"""
        try:
            print(f"📝 DOCX dosyası işleniyor: {file_path}")
            doc = DocxDocument(file_path)
            text = ""
            
            print(f"📊 DOCX paragraf sayısı: {len(doc.paragraphs)}")
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    if i < 5:  # İlk 5 paragrafı logla
                        print(f"📝 Paragraf {i+1}: {paragraph.text[:100]}...")
            
            # Metni temizle ve normalize et
            cleaned_text = self._clean_and_normalize_text(text)
            print(f"✅ DOCX metin çıkarma başarılı: {len(cleaned_text)} karakter")
            return cleaned_text
            
        except Exception as e:
            print(f"❌ DOCX extraction error: {e}")
            print(f"❌ Hata türü: {type(e)}")
            return ""
    
    async def _extract_doc_text(self, file_path: str) -> str:
        """DOC dosyasından metin çıkar - Eski Word formatı"""
        try:
            print(f"📝 DOC dosyası işleniyor: {file_path}")
            
            # Dosya header'ını kontrol et
            with open(file_path, 'rb') as f:
                header = f.read(16)
                print(f"🔍 DOC header (hex): {header.hex()}")
                print(f"🔍 DOC header (ascii): {repr(header)}")
            
            # 1. Yöntem: python-docx ile dene (bazen .doc dosyaları .docx olarak açılabilir)
            try:
                print("📝 python-docx ile DOC deneniyor...")
                doc = DocxDocument(file_path)
                text = ""
                
                for i, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                        if i < 5:
                            print(f"📝 Paragraf {i+1}: {paragraph.text[:100]}...")
                
                if text.strip():
                    print(f"✅ python-docx ile DOC başarılı: {len(text)} karakter")
                    return text.strip()
                else:
                    print("⚠️ python-docx ile çıkarılan metin boş")
                    
            except Exception as docx_error:
                print(f"❌ python-docx ile DOC hatası: {docx_error}")
            
            # 2. Yöntem: Windows COM ile dene (Windows sistemlerde)
            try:
                print("📝 Windows COM ile DOC deneniyor...")
                import win32com.client
                import os
                
                # Word uygulamasını başlat
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # Dökümanı aç
                doc = word.Documents.Open(os.path.abspath(file_path))
                text = doc.Content.Text
                
                # Dökümanı kapat
                doc.Close()
                word.Quit()
                
                if text.strip():
                    print(f"✅ Windows COM ile DOC başarılı: {len(text)} karakter")
                    return text.strip()
                else:
                    print("⚠️ Windows COM ile çıkarılan metin boş")
                    
            except Exception as com_error:
                print(f"❌ Windows COM hatası: {com_error}")
            
            # 3. Son çare: Düz metin olarak oku
            try:
                print("📝 Düz metin olarak DOC deneniyor...")
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                
                if len(text.strip()) > 10:
                    print(f"✅ Düz metin olarak DOC başarılı: {len(text)} karakter")
                    return text.strip()
                else:
                    print("⚠️ Düz metin çok kısa")
                    
            except Exception as text_error:
                print(f"❌ Düz metin hatası: {text_error}")
            
            print("❌ DOC dosyası hiçbir yöntemle okunamadı")
            return "DOC dosyasından metin çıkarılamadı: Desteklenmeyen format"
            
        except Exception as e:
            print(f"❌ DOC metin çıkarma genel hatası: {e}")
            print(f"❌ Hata türü: {type(e)}")
            return f"DOC dosyasından metin çıkarılamadı: {str(e)}"
    
    async def _extract_excel_text(self, file_path: str) -> str:
        """Excel dosyasından metin çıkar - Hem .xlsx hem .xls formatlarını destekler"""
        try:
            text = ""
            
            # Dosya detaylarını logla
            print(f"🔍 Excel dosya analizi başlıyor...")
            print(f"📁 Dosya yolu: {file_path}")
            print(f"📏 Dosya boyutu: {os.path.getsize(file_path)} bytes")
            
            # Dosya header'ını kontrol et
            with open(file_path, 'rb') as f:
                header = f.read(16)  # İlk 16 byte
                print(f"🔍 Dosya header (hex): {header.hex()}")
                print(f"🔍 Dosya header (ascii): {repr(header)}")
            
            # Dosya uzantısına göre uygun yöntemi kullan
            if file_path.lower().endswith('.xlsx'):
                print("📊 .xlsx dosyası openpyxl ile işleniyor...")
                try:
                    workbook = openpyxl.load_workbook(file_path, data_only=True)
                    
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        text += f"\n--- {sheet_name} ---\n"
                        
                        for row in sheet.iter_rows(values_only=True):
                            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                            if row_text.strip():
                                text += row_text + "\n"
                    
                    workbook.close()
                    print(f"✅ openpyxl başarılı: {len(text)} karakter")
                    
                except Exception as xlsx_error:
                    print(f"❌ openpyxl hatası: {xlsx_error}")
                    raise xlsx_error
                
            elif file_path.lower().endswith('.xls'):
                print("📊 .xls dosyası xlrd ile işleniyor...")
                try:
                    import xlrd
                    
                    # xlrd versiyonunu kontrol et
                    print(f"📊 xlrd versiyonu: {xlrd.__version__}")
                    
                    # Dosya formatını kontrol et
                    try:
                        workbook = xlrd.open_workbook(file_path)
                        print(f"✅ xlrd workbook açıldı, sheet sayısı: {workbook.nsheets}")
                        
                        for sheet_name in workbook.sheet_names():
                            sheet = workbook.sheet_by_name(sheet_name)
                            text += f"\n--- {sheet_name} ---\n"
                            print(f"📊 Sheet işleniyor: {sheet_name} ({sheet.nrows} satır, {sheet.ncols} sütun)")
                            
                            for row_idx in range(sheet.nrows):
                                row_values = sheet.row_values(row_idx)
                                row_text = " | ".join([str(cell) if cell != "" else "" for cell in row_values])
                                if row_text.strip():
                                    text += row_text + "\n"
                        
                        print(f"✅ xlrd başarılı: {len(text)} karakter")
                        
                    except Exception as xlrd_error:
                        print(f"❌ xlrd hatası: {xlrd_error}")
                        print(f"❌ Hata türü: {type(xlrd_error)}")
                        
                        # Alternatif yöntemler dene
                        print("🔄 Alternatif yöntemler deneniyor...")
                        
                        # 1. pandas ile dene
                        try:
                            import pandas as pd
                            print("📊 pandas ile deneniyor...")
                            df = pd.read_excel(file_path, engine='xlrd')
                            text = f"Pandas ile okunan veri:\n{df.to_string()}"
                            print(f"✅ pandas başarılı: {len(text)} karakter")
                            return text.strip()
                        except Exception as pd_error:
                            print(f"❌ pandas hatası: {pd_error}")
                        
                        # 2. openpyxl ile dene (bazen .xls dosyaları .xlsx olarak açılabilir)
                        try:
                            print("📊 openpyxl ile .xls deneniyor...")
                            workbook = openpyxl.load_workbook(file_path, data_only=True)
                            for sheet_name in workbook.sheetnames:
                                sheet = workbook[sheet_name]
                                text += f"\n--- {sheet_name} ---\n"
                                for row in sheet.iter_rows(values_only=True):
                                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                                    if row_text.strip():
                                        text += row_text + "\n"
                            workbook.close()
                            print(f"✅ openpyxl .xls başarılı: {len(text)} karakter")
                            return text.strip()
                        except Exception as openpyxl_xls_error:
                            print(f"❌ openpyxl .xls hatası: {openpyxl_xls_error}")
                        
                        raise xlrd_error
                        
                except Exception as xlrd_error:
                    print(f"❌ xlrd genel hatası: {xlrd_error}")
                    raise xlrd_error
                
            else:
                print(f"⚠️ Bilinmeyen Excel formatı: {file_path}")
                return f"Desteklenmeyen Excel formatı: {file_path}"
            
            if not text.strip():
                print("⚠️ Çıkarılan metin boş!")
                return "Excel dosyasından metin çıkarılamadı: Boş içerik"
            
            print(f"✅ Excel metin çıkarma başarılı: {len(text)} karakter")
            return text.strip()
            
        except Exception as e:
            print(f"❌ Excel metin çıkarma hatası: {e}")
            print(f"❌ Hata türü: {type(e)}")
            print(f"❌ Hata detayı: {str(e)}")
            
            # Dosya hakkında ek bilgi
            try:
                import magic
                mime_type = magic.from_file(file_path, mime=True)
                print(f"🔍 MIME type: {mime_type}")
            except:
                print("🔍 MIME type tespit edilemedi")
            
            return f"Excel dosyasından metin çıkarılamadı: {str(e)}"
    
    async def _extract_image_text(self, file_path: str) -> str:
        """Resimden OCR ile metin çıkar - Gemini Vision API kullanarak"""
        try:
            print(f"🖼️ Gemini Vision API ile resim işleniyor: {file_path}")
            
            # Gemini Vision API kullan
            extracted_text = await self.gemini_service.extract_text_from_image(file_path)
            
            if extracted_text and len(extracted_text.strip()) > 10:
                print(f"✅ Gemini OCR başarılı: {len(extracted_text)} karakter çıkarıldı")
                return extracted_text.strip()
            else:
                print("⚠️ Gemini OCR sonucu çok kısa, Tesseract fallback kullanılıyor...")
                # Fallback olarak Tesseract kullan
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image, lang='tur+eng')
                print(f"✅ Tesseract fallback başarılı: {len(text)} karakter çıkarıldı")
                return text.strip()
                
        except Exception as e:
            print(f"❌ Gemini OCR hatası: {e}")
            try:
                print("🔄 Tesseract fallback deneniyor...")
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image, lang='tur+eng')
                print(f"✅ Tesseract fallback başarılı: {len(text)} karakter çıkarıldı")
                return text.strip()
            except Exception as fallback_error:
                print(f"❌ Tesseract fallback da başarısız: {fallback_error}")
                return f"Resimden metin çıkarılamadı. Hata: {str(e)}"
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """TXT dosyasından metin çıkar"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    async def _detect_real_file_type(self, file_path: str, file_extension: str) -> str:
        """Dosya header'ına göre gerçek formatı tespit et"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(16)  # İlk 16 byte
            
            print(f"🔍 Dosya header analizi: {header.hex()}")
            
            # PDF tespiti
            if header.startswith(b'%PDF'):
                print("📄 PDF formatı tespit edildi")
                return "pdf"
            
            # DOCX tespiti (ZIP formatı)
            elif header.startswith(b'PK\x03\x04'):
                print("📝 DOCX formatı tespit edildi")
                return "docx"
            
            # DOC tespiti (OLE formatı)
            elif header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                print("📝 DOC formatı tespit edildi")
                return "doc"
            
            # XLSX tespiti (ZIP formatı)
            elif header.startswith(b'PK\x03\x04'):
                print("📊 XLSX formatı tespit edildi")
                return "xlsx"
            
            # XLS tespiti (OLE formatı)
            elif header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                print("📊 XLS formatı tespit edildi")
                return "xls"
            
            # HTML tespiti
            elif header.startswith(b'<html') or header.startswith(b'<!DOCTYPE') or header.startswith(b'<'):
                print("🌐 HTML formatı tespit edildi")
                return "html"
            
            # JPEG tespiti
            elif header.startswith(b'\xff\xd8\xff'):
                print("🖼️ JPEG formatı tespit edildi")
                return "jpeg"
            
            # PNG tespiti
            elif header.startswith(b'\x89PNG\r\n\x1a\n'):
                print("🖼️ PNG formatı tespit edildi")
                return "png"
            
            # TXT tespiti (basit metin)
            elif all(32 <= b <= 126 or b in [9, 10, 13] for b in header[:10]):
                print("📄 TXT formatı tespit edildi")
                return "txt"
            
            else:
                print(f"⚠️ Bilinmeyen format, uzantıya göre işleniyor: {file_extension}")
                return file_extension.lower()
                
        except Exception as e:
            print(f"❌ Format tespit hatası: {e}")
            return file_extension.lower()
    
    async def _extract_html_text(self, file_path: str) -> str:
        """HTML dosyasından metin çıkar - Tag'leri temizleyerek"""
        try:
            print(f"🌐 HTML dosyası işleniyor: {file_path}")
            
            # Dosya header'ını kontrol et
            with open(file_path, 'rb') as f:
                header = f.read(50)  # İlk 50 byte
                print(f"🔍 HTML header: {repr(header)}")
            
            # HTML'i oku
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Basit HTML tag temizleme
            import re
            
            # Script ve style tag'lerini kaldır
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            
            # HTML tag'lerini kaldır
            html_content = re.sub(r'<[^>]+>', '', html_content)
            
            # Çoklu boşlukları temizle
            html_content = re.sub(r'\s+', ' ', html_content)
            
            # HTML entities'leri decode et
            html_content = html_content.replace('&nbsp;', ' ')
            html_content = html_content.replace('&amp;', '&')
            html_content = html_content.replace('&lt;', '<')
            html_content = html_content.replace('&gt;', '>')
            html_content = html_content.replace('&quot;', '"')
            
            # Başındaki ve sonundaki boşlukları temizle
            html_content = html_content.strip()
            
            print(f"✅ HTML metin çıkarma başarılı: {len(html_content)} karakter")
            return html_content
            
        except Exception as e:
            print(f"❌ HTML metin çıkarma hatası: {e}")
            # Fallback olarak düz metin olarak oku
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    raw_content = file.read()
                print(f"✅ HTML fallback başarılı: {len(raw_content)} karakter")
                return raw_content
            except Exception as fallback_error:
                print(f"❌ HTML fallback da başarısız: {fallback_error}")
                return f"HTML dosyasından metin çıkarılamadı: {str(e)}"
    
    async def create_document_chunks(self, document: Document, content_text: str, db: Session):
        """Dökümanı parçalara böl ve kaydet - Optimize edilmiş"""
        try:
            print(f"✂️ Chunk oluşturma başlıyor...")
            print(f"📏 Toplam metin uzunluğu: {len(content_text)} karakter")
            
            chunks = self._split_text_into_chunks(content_text, chunk_size=1200, overlap=250)
            print(f"📊 Toplam {len(chunks)} chunk oluşturuldu")
            
            # Batch processing için chunks'ları grupla
            batch_size = 5
            total_batches = (len(chunks) + batch_size - 1) // batch_size
            print(f"🔄 {total_batches} batch işlenecek")
            
            for batch_num, i in enumerate(range(0, len(chunks), batch_size), 1):
                batch_chunks = chunks[i:i+batch_size]
                print(f"🔄 Batch {batch_num}/{total_batches} işleniyor ({len(batch_chunks)} chunk)")
                
                for j, chunk_text in enumerate(batch_chunks):
                    chunk_index = i + j
                    print(f"📝 Chunk {chunk_index + 1}/{len(chunks)} işleniyor ({len(chunk_text)} karakter)")
                    
                    try:
                        # Her parça için embedding oluştur
                        print(f"🧠 Chunk {chunk_index + 1} için embedding oluşturuluyor...")
                        chunk_embeddings = await self.gemini_service.generate_embeddings(chunk_text)
                        print(f"✅ Chunk {chunk_index + 1} embedding tamamlandı: {len(chunk_embeddings)} vektör")
                        
                        chunk = DocumentChunk(
                            document_id=document.id,
                            chunk_text=chunk_text,
                            chunk_index=chunk_index,
                            embeddings=json.dumps(chunk_embeddings)
                        )
                        db.add(chunk)
                        print(f"✅ Chunk {chunk_index + 1} veritabanına eklendi")
                        
                    except Exception as chunk_error:
                        print(f"❌ Chunk {chunk_index + 1} hatası: {chunk_error}")
                        # Hata durumunda boş embedding kullan
                        chunk = DocumentChunk(
                            document_id=document.id,
                            chunk_text=chunk_text,
                            chunk_index=chunk_index,
                            embeddings=json.dumps([])
                        )
                        db.add(chunk)
                        print(f"⚠️ Chunk {chunk_index + 1} boş embedding ile eklendi")
                
                # Batch'i commit et
                print(f"💾 Batch {batch_num} commit ediliyor...")
                db.commit()
                print(f"✅ Batch {batch_num} commit edildi")
            
            print(f"🎉 Tüm chunk'lar başarıyla oluşturuldu ve kaydedildi")
            
        except Exception as e:
            print(f"❌ Chunk oluşturma hatası: {e}")
            print(f"❌ Hata türü: {type(e)}")
            import traceback
            print(f"❌ Traceback: {traceback.format_exc()}")
            raise e
    
    def _clean_and_normalize_text(self, text: str) -> str:
        """Metni temizle ve normalize et - PDF'den gelen parçalı metni düzelt"""
        if not text:
            return ""
        
        # Gereksiz boşlukları ve satır sonlarını temizle
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Boş satırları atla
                # Satır sonundaki tire işaretlerini kaldır (kelime bölünmesi)
                if line.endswith('-') and len(line) > 1:
                    line = line[:-1]
                cleaned_lines.append(line)
        
        # Satırları birleştir ve normal cümle yapısına getir
        normalized_text = ' '.join(cleaned_lines)
        
        # Çoklu boşlukları tek boşluk yap
        import re
        normalized_text = re.sub(r'\s+', ' ', normalized_text)
        
        # Noktalama işaretlerinden sonra tek boşluk bırak
        normalized_text = re.sub(r'([.!?])\s*', r'\1 ', normalized_text)
        
        # Paragraf yapısını koru (çift satır sonu)
        normalized_text = re.sub(r'\n\s*\n', '\n\n', normalized_text)
        
        print(f"🧹 Text cleaned: {len(text)} → {len(normalized_text)} characters")
        return normalized_text.strip()
    
    def _split_text_into_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
        """Metni overlapping parçalara böl"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Kelime ortasında kesmemek için
            if end < len(text) and not text[end].isspace():
                last_space = chunk.rfind(' ')
                if last_space > 0:
                    chunk = chunk[:last_space]
                    end = start + last_space
            
            chunks.append(chunk.strip())
            start = end - overlap
            
            if start >= len(text):
                break
        
        return [chunk for chunk in chunks if chunk.strip()]
